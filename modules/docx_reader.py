"""
DOCX 内容提取模块（Phase 4 Step 1: Reader）
负责从 .docx 文件中提取文字、标题、表格和图片。

设计原则：
  - 完全独立，不依赖 PDF 模块
  - 不调用 OCR（图片只保存，建立映射）
  - 不调用翻译
  - 复用 modules/ 的 TextBlock / ImageBlock 数据结构（仅复用定义，不依赖功能）

数据结构：
  DocxContent: DOCX 文档的完整中间表示
  DocxParagraph: 普通段落（文字 + 样式信息）
  DocxHeading: 标题（级别 + 文字）
  DocxTable: 表格（二维单元格）
  DocxImage: 图片（保存到磁盘 + 映射关系）
"""

import os
import sys
import io
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

from docx import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image


# ── DOCX 专用数据结构 ──────────────────────────────────────

@dataclass
class DocxParagraph:
    """普通段落"""
    text: str
    style_name: str = ""            # 段落样式名（如 "Normal"）
    font_name: str = ""             # 字体名
    font_size: float = 10.0         # 字号（pt）
    bold: bool = False
    italic: bool = False
    alignment: str = "left"         # left / center / right / justify
    first_line_indent: float = 0.0  # 首行缩进（pt）
    line_spacing: float = 1.15      # 行距
    color: str = "#000000"          # 文字颜色
    index: int = 0                  # 在文档中的序号（用于保持顺序）
    run_count: int = 0              # 该段落包含的 run 数量


@dataclass
class DocxHeading:
    """标题"""
    text: str
    level: int                      # 1-9（对应 H1-H9）
    style_name: str = ""
    font_name: str = ""
    font_size: float = 16.0
    bold: bool = True
    alignment: str = "left"
    index: int = 0


@dataclass
class DocxTableCell:
    """表格单元格"""
    text: str
    row: int
    col: int
    font_name: str = ""
    font_size: float = 10.0
    bold: bool = False
    alignment: str = "left"


@dataclass
class DocxTable:
    """表格"""
    rows: int
    cols: int
    cells: List[DocxTableCell] = field(default_factory=list)
    style_name: str = ""
    index: int = 0

    def get_cell(self, row: int, col: int) -> Optional[DocxTableCell]:
        for cell in self.cells:
            if cell.row == row and cell.col == col:
                return cell
        return None


@dataclass
class DocxImage:
    """图片"""
    image: Optional[Image.Image] = None   # PIL Image 对象
    image_path: str = ""                  # 保存到磁盘的路径
    ext: str = "png"                      # 扩展名
    width_px: int = 0
    height_px: int = 0
    paragraph_index: int = 0              # 所在段落序号
    content_type: str = ""                # MIME 类型
    index: int = 0                        # 图片序号

    @property
    def size(self) -> Tuple[int, int]:
        return (self.width_px, self.height_px)


@dataclass
class DocxContent:
    """DOCX 文档的完整提取结果"""
    file_path: str
    paragraphs: List[DocxParagraph] = field(default_factory=list)
    headings: List[DocxHeading] = field(default_factory=list)
    tables: List[DocxTable] = field(default_factory=list)
    images: List[DocxImage] = field(default_factory=list)
    headers: List[str] = field(default_factory=list)   # 页眉文字列表
    footers: List[str] = field(default_factory=list)   # 页脚文字列表

    @property
    def total_text_count(self) -> int:
        return len(self.paragraphs) + len(self.headings)

    @property
    def total_table_count(self) -> int:
        return len(self.tables)

    @property
    def total_image_count(self) -> int:
        return len(self.images)


# ── DOCX Reader ────────────────────────────────────────────

class DocxReader:
    """
    DOCX 读取器。

    使用方式:
        reader = DocxReader("path/to/file.docx", temp_dir="./temp/docx")
        content = reader.extract()

        print(f"段落: {len(content.paragraphs)}")
        print(f"标题: {len(content.headings)}")
        print(f"表格: {len(content.tables)}")
        print(f"图片: {len(content.images)}")
    """

    def __init__(self, file_path: str, temp_dir: str = None):
        self.file_path = file_path
        self.file_name = Path(file_path).stem

        # 临时目录（保存提取的图片）
        if temp_dir is None:
            temp_dir = os.path.join(
                os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                "temp", "docx", self.file_name
            )
        self.temp_dir = temp_dir
        os.makedirs(self.temp_dir, exist_ok=True)

        # 加载文档
        self.doc = DocxDocument(file_path)

    def extract(self) -> DocxContent:
        """提取 DOCX 全部内容"""
        content = DocxContent(file_path=self.file_path)

        # 提取顺序：段落/标题（交替出现）→ 表格 → 图片 → 页眉页脚
        content.paragraphs = self._extract_paragraphs()
        content.headings = self._extract_headings()
        content.tables = self._extract_tables()
        content.images = self._extract_images()
        content.headers = self._extract_headers_footers("header")
        content.footers = self._extract_headers_footers("footer")

        return content

    # ── 段落提取 ───────────────────────────────────────────

    def _extract_paragraphs(self) -> List[DocxParagraph]:
        """提取所有普通段落（非标题、非空）"""
        paragraphs = []
        para_index = 0

        for para in self.doc.paragraphs:
            # 跳过标题样式段落（由 _extract_headings 单独处理）
            if para.style.name and para.style.name.startswith("Heading"):
                para_index += 1
                continue

            text = para.text.strip()
            if not text:
                para_index += 1
                continue

            # 提取段落样式信息
            style_info = self._parse_paragraph_style(para)

            paragraphs.append(DocxParagraph(
                text=text,
                style_name=para.style.name if para.style else "",
                font_name=style_info["font_name"],
                font_size=style_info["font_size"],
                bold=style_info["bold"],
                italic=style_info["italic"],
                alignment=style_info["alignment"],
                first_line_indent=style_info["first_line_indent"],
                line_spacing=style_info["line_spacing"],
                color=style_info["color"],
                index=para_index,
                run_count=len(para.runs),
            ))
            para_index += 1

        return paragraphs

    # ── 标题提取 ───────────────────────────────────────────

    def _extract_headings(self) -> List[DocxHeading]:
        """提取所有标题（Heading 1-9）"""
        headings = []
        para_index = 0

        for para in self.doc.paragraphs:
            style_name = para.style.name if para.style else ""

            # 只处理标题样式
            if not style_name.startswith("Heading"):
                para_index += 1
                continue

            text = para.text.strip()
            if not text:
                para_index += 1
                continue

            # 解析标题级别
            level_str = style_name.replace("Heading", "").strip()
            try:
                level = int(level_str)
            except ValueError:
                level = 1

            style_info = self._parse_paragraph_style(para)

            headings.append(DocxHeading(
                text=text,
                level=level,
                style_name=style_name,
                font_name=style_info["font_name"],
                font_size=style_info["font_size"],
                bold=style_info["bold"],
                alignment=style_info["alignment"],
                index=para_index,
            ))
            para_index += 1

        return headings

    # ── 表格提取 ───────────────────────────────────────────

    def _extract_tables(self) -> List[DocxTable]:
        """提取所有表格"""
        tables = []

        for table_idx, table in enumerate(self.doc.tables):
            cells = []
            max_cols = 0

            for row_idx, row in enumerate(table.rows):
                col_count = len(row.cells)
                max_cols = max(max_cols, col_count)

                for col_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue

                    # 提取单元格样式
                    font_name, font_size, bold, alignment = self._parse_cell_style(cell)

                    cells.append(DocxTableCell(
                        text=text,
                        row=row_idx,
                        col=col_idx,
                        font_name=font_name,
                        font_size=font_size,
                        bold=bold,
                        alignment=alignment,
                    ))

            if cells:
                tables.append(DocxTable(
                    rows=len(table.rows),
                    cols=max_cols,
                    cells=cells,
                    style_name=table.style.name if table.style else "",
                    index=table_idx,
                ))

        return tables

    # ── 图片提取 ───────────────────────────────────────────

    def _extract_images(self) -> List[DocxImage]:
        """提取文档中所有内嵌图片（只保存，不 OCR）"""
        images = []
        img_index = 0

        # 遍历所有段落和内联图片
        for para_idx, para in enumerate(self.doc.paragraphs):
            # 通过 XML 查找图片关系
            drawings = para._element.findall(
                './/' + qn('wp:inline')  # 内联图片
            ) + para._element.findall(
                './/' + qn('wp:anchor')  # 浮动图片
            )

            for drawing in drawings:
                # 查找图片引用
                blip_elements = drawing.findall('.//' + qn('a:blip'))
                for blip in blip_elements:
                    embed_id = blip.get(qn('r:embed'))
                    if not embed_id:
                        continue

                    try:
                        # 获取图片数据
                        image_part = self.doc.part.related_parts[embed_id]
                        image_bytes = image_part.blob
                        content_type = image_part.content_type

                        # 确定扩展名
                        ext = self._content_type_to_ext(content_type)

                        # 保存到磁盘
                        img_filename = f"image_{img_index + 1}.{ext}"
                        img_path = os.path.join(self.temp_dir, img_filename)
                        with open(img_path, "wb") as f:
                            f.write(image_bytes)

                        # 尝试转换为 PIL Image
                        pil_image = None
                        width_px, height_px = 0, 0
                        try:
                            pil_image = Image.open(io.BytesIO(image_bytes))
                            width_px, height_px = pil_image.size
                        except Exception:
                            pass

                        # 尝试从 XML 获取尺寸信息
                        extents = drawing.findall('.//' + qn('wp:extent'))
                        if extents:
                            cx = int(extents[0].get('cx', 0))   # EMU
                            cy = int(extents[0].get('cy', 0))
                            if cx > 0 and cy > 0 and width_px == 0:
                                # EMU → px (1 EMU = 1/914400 inch, 1 inch ≈ 96px)
                                width_px = int(cx / 914400 * 96)
                                height_px = int(cy / 914400 * 96)

                        images.append(DocxImage(
                            image=pil_image,
                            image_path=img_path,
                            ext=ext,
                            width_px=width_px,
                            height_px=height_px,
                            paragraph_index=para_idx,
                            content_type=content_type,
                            index=img_index,
                        ))
                        img_index += 1

                    except KeyError:
                        # 某些内嵌资源可能无法访问，跳过
                        continue
                    except Exception:
                        continue

        return images

    # ── 页眉页脚 ───────────────────────────────────────────

    def _extract_headers_footers(self, section_type: str) -> List[str]:
        """提取页眉或页脚文字"""
        texts = []
        for section in self.doc.sections:
            if section_type == "header":
                hf = section.header
            else:
                hf = section.footer

            if hf and not hf.is_linked_to_previous:
                for para in hf.paragraphs:
                    if para.text.strip():
                        texts.append(para.text.strip())
        return texts

    # ── 样式解析 ───────────────────────────────────────────

    def _parse_paragraph_style(self, para) -> dict:
        """解析段落样式信息"""
        result = {
            "font_name": "",
            "font_size": 10.0,
            "bold": False,
            "italic": False,
            "alignment": "left",
            "first_line_indent": 0.0,
            "line_spacing": 1.15,
            "color": "#000000",
        }

        # 对齐方式
        if para.alignment is not None:
            align_map = {
                WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
                WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
                WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
                WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify",
            }
            result["alignment"] = align_map.get(para.alignment, "left")

        # 段落格式
        pf = para.paragraph_format
        if pf:
            if pf.first_line_indent:
                result["first_line_indent"] = pf.first_line_indent.pt
            if pf.line_spacing:
                result["line_spacing"] = pf.line_spacing

        # 获取第一个 run 的字体信息
        if para.runs:
            run = para.runs[0]
            result["font_name"] = run.font.name or ""
            if run.font.size:
                result["font_size"] = run.font.size.pt
            result["bold"] = run.bold or False
            result["italic"] = run.italic or False
            if run.font.color and run.font.color.rgb:
                result["color"] = "#" + str(run.font.color.rgb)

        return result

    def _parse_cell_style(self, cell) -> tuple:
        """解析单元格样式，返回 (font_name, font_size, bold, alignment)"""
        font_name, font_size, bold, alignment = "", 10.0, False, "left"

        for para in cell.paragraphs:
            if para.runs:
                run = para.runs[0]
                font_name = run.font.name or font_name
                if run.font.size:
                    font_size = run.font.size.pt
                bold = run.bold or bold
                if para.alignment is not None:
                    align_map = {
                        WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
                        WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
                        WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
                    }
                    alignment = align_map.get(para.alignment, "left")
                break

        return font_name, font_size, bold, alignment

    def _content_type_to_ext(self, content_type: str) -> str:
        """MIME 类型 → 文件扩展名"""
        mapping = {
            "image/png": "png",
            "image/jpeg": "jpg",
            "image/gif": "gif",
            "image/bmp": "bmp",
            "image/tiff": "tiff",
            "image/webp": "webp",
            "image/x-emf": "emf",
            "image/x-wmf": "wmf",
        }
        return mapping.get(content_type, "png")

    # ── 摘要 ───────────────────────────────────────────────

    def summary(self, content: DocxContent = None) -> str:
        """生成提取结果摘要"""
        if content is None:
            content = self.extract()

        lines = [
            f"文件: {self.file_name}.docx",
            f"段落: {len(content.paragraphs)} 个",
            f"标题: {len(content.headings)} 个",
            f"表格: {len(content.tables)} 个 ({sum(t.rows for t in content.tables)} 行)",
            f"图片: {len(content.images)} 个",
            f"页眉: {len(content.headers)} 处",
            f"页脚: {len(content.footers)} 处",
        ]

        if content.headings:
            lines.append("\n标题结构:")
            for h in sorted(content.headings, key=lambda x: x.index):
                indent = "  " * (h.level - 1)
                lines.append(f"  {indent}H{h.level}: {h.text[:60]}")

        return "\n".join(lines)
